"""
tests/test_local_pipeline.py — Unit Test Suite for Sovereign Workbench Local Modules
SIH26117 / MoPNG / MRPL Agentic AI Workbench (v5.3 Architecture)

Coverage:
  1. src.router       — L1 fast-path (<5ms), mime/extension routing, calculation regex, L2 fallback.
  2. src.density_gate — Blank/whitespace guard, short-string guard, null-byte/control chars, valid text.
  3. src.loop_killer  — Deterministic sort_keys hashing, error tail differentiation.
  4. src.exporter     — Run-flattening placeholder substitution, empty paragraph resilience, citation contract.
"""

import os
import time
from unittest.mock import MagicMock
import pytest
from docx import Document

from src.router import route_l1, route_l2
from src.density_gate import is_gibberish_or_blank
from src.loop_killer import get_step_hash, hash_observation_tail
from src.exporter import render_deliverable


# ===========================================================================
# 1. Tests for src.router
# ===========================================================================

class TestRouter:
    """Tests for deterministic L1 routing and L2 fallback handling."""

    def test_l1_image_mime_and_extensions_to_vision(self):
        """Verify .png/.jpg and image/* MIME route to vision in < 5 ms."""
        test_cases = [
            (["image/png"], ["diagram.png"]),
            (["image/jpeg"], ["scan.jpg"]),
            (["image/tiff"], ["p_and_id.tiff"]),
            (["image/webp"], ["unit200.webp"]),
        ]

        for mimes, names in test_cases:
            start = time.perf_counter()
            specialist, trace = route_l1(mimes, names, "Inspect P&ID tag", "clean")
            elapsed_ms = (time.perf_counter() - start) * 1000

            assert specialist == "vision", f"Expected vision for {names}, got {specialist}"
            assert "Vision" in trace
            assert elapsed_ms < 5.0, f"L1 routing exceeded latency budget: {elapsed_ms:.2f} ms"

    def test_l1_spreadsheet_extensions_to_coder(self):
        """Verify .xlsx/.csv files route to coder in < 5 ms."""
        test_cases = [
            (["application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"], ["measurements.xlsx"]),
            (["text/csv"], ["inspection_data.csv"]),
            (["application/octet-stream"], ["readings.XLSX"]),
            (["application/octet-stream"], ["readings.CSV"]),
        ]

        for mimes, names in test_cases:
            start = time.perf_counter()
            specialist, trace = route_l1(mimes, names, "Load data table", "clean")
            elapsed_ms = (time.perf_counter() - start) * 1000

            assert specialist == "coder", f"Expected coder for {names}, got {specialist}"
            assert "Coder" in trace
            assert elapsed_ms < 5.0, f"L1 routing exceeded latency budget: {elapsed_ms:.2f} ms"

    def test_l1_calculation_keywords_to_coder(self):
        """Verify calculation keywords (calculate, compute, trend, rate, delta) route to coder."""
        calc_prompts = [
            "Please calculate corrosion rate for pipe P-101",
            "Compute the average thickness over 12 months",
            "What is the historical trend of wall degradation?",
            "Assess the maximum allowable rate of erosion",
            "Find the delta between baseline and Q3 inspection",
            "Calculate and plot trend for Unit 200",
        ]

        for prompt in calc_prompts:
            specialist, trace = route_l1([], [], prompt, "clean")
            assert specialist == "coder", f"Failed to match calc keyword in: '{prompt}'"
            assert "math-intent" in trace

    def test_l1_density_outcome_to_vision(self):
        """Verify blank, scanned, or gibberish page outcome forces Vision OCR."""
        for outcome in ("blank", "scanned", "gibberish"):
            specialist, trace = route_l1([], ["document.pdf"], "Summarize notes", outcome)
            assert specialist == "vision"
            assert "Vision OCR" in trace

    def test_l1_unhandled_prompt_falls_through_to_l2(self):
        """Verify unhandled general queries fall through to L2 without raising exceptions."""
        specialist, trace = route_l1(
            mimes=["application/pdf"],
            names=["general_sop.pdf"],
            prompt="Summarize section 4 safety guidelines for refinery personnel",
            page_outcome="clean",
        )
        assert specialist is None
        assert "L2" in trace

    def test_l2_judge_execution_mocked(self):
        """Verify route_l2 safely queries Brain-7B constrained judge."""
        mock_brain = MagicMock()
        mock_brain.chat.return_value = {
            "specialist": "rag",
            "confidence": 0.98,
        }

        specialist, trace = route_l2("Summarize unit operating standard", mock_brain)
        assert specialist == "rag"
        assert "L2 Brain-7B judge -> rag" in trace
        mock_brain.chat.assert_called_once()


# ===========================================================================
# 2. Tests for src.density_gate
# ===========================================================================

class TestDensityGate:
    """Tests for math-safe density gate and gibberish filtration."""

    def test_empty_and_whitespace_only(self):
        """Empty or whitespace-only strings must return True (zero-division safe)."""
        assert is_gibberish_or_blank("") is True
        assert is_gibberish_or_blank("   ") is True
        assert is_gibberish_or_blank("\n\t  \r\n") is True
        assert is_gibberish_or_blank(None) is True

    def test_short_strings_under_50_chars(self):
        """Strings shorter than 50 characters cannot yield reliable ratio and must return True."""
        assert is_gibberish_or_blank("Short line") is True
        assert is_gibberish_or_blank("A" * 49) is True
        assert is_gibberish_or_blank("   " + "B" * 48 + "   ") is True

    def test_null_byte_and_control_gibberish(self):
        """Strings containing null bytes or binary control characters must return True."""
        prefix = "Valid engineering document text that has more than 50 characters."
        assert is_gibberish_or_blank(prefix + "\x00" + "suffix text") is True
        assert is_gibberish_or_blank(prefix + "\x07" + "bell character") is True
        assert is_gibberish_or_blank(prefix + "\x1f" + "unit separator") is True
        assert is_gibberish_or_blank(prefix + "\x7f" + "delete control") is True

    def test_low_alphanumeric_ratio(self):
        """Strings with alpha/space ratio < 0.40 must return True."""
        # Mostly punctuation/symbols
        symbol_text = "!@#$%^&*()_+=~`{}[]|:;'<>?,./ " * 5
        assert len(symbol_text) >= 50
        assert is_gibberish_or_blank(symbol_text) is True

    def test_valid_engineering_text(self):
        """Valid English narrative text with len >= 50 and ratio >= 0.40 must return False."""
        valid_doc = (
            "Mangalore Refinery and Petrochemicals Limited (MRPL). "
            "Operating manual for Crude Distillation Unit 200. "
            "All safety interlocks must be verified before commencing startup."
        )
        assert len(valid_doc) >= 50
        assert is_gibberish_or_blank(valid_doc) is False


# ===========================================================================
# 3. Tests for src.loop_killer
# ===========================================================================

class TestLoopKiller:
    """Tests for deterministic observation tail hashing and loop suppression."""

    def test_deterministic_key_sorting(self):
        """Assert get_step_hash produces identical hashes regardless of args dict key insertion order."""
        args_order_1 = {"task_id": "T-100", "step": 3, "query": "Unit 200", "n_results": 5}
        args_order_2 = {"n_results": 5, "query": "Unit 200", "step": 3, "task_id": "T-100"}

        hash1 = get_step_hash("rag_search", args_order_1, "Sample observation tail")
        hash2 = get_step_hash("rag_search", args_order_2, "Sample observation tail")

        assert hash1 == hash2
        assert len(hash1) == 64  # Valid SHA-256 hex digest

    def test_alias_equivalence(self):
        """Verify get_step_hash and hash_observation_tail produce identical outputs."""
        args = {"tool": "sandbox_run", "limit": 100}
        obs = "Output traceback tail"
        assert get_step_hash("sandbox_run", args, obs) == hash_observation_tail("sandbox_run", args, obs)

    def test_differing_error_tails_produce_distinct_hashes(self):
        """Tracebacks sharing identical 1000-character headers must produce distinct hashes if tails differ."""
        identical_header = "Traceback (most recent call last):\n" + (
            "  File '/srv/sovereign/sandbox.py', line 120, in execute_calc\n"
            "    result = compute_thickness(readings, baseline)\n"
        ) * 15  # Well over 500 chars

        tail_error_a = "ZeroDivisionError: division by zero in rate calculation"
        tail_error_b = "KeyError: 'temperature_celsius' column missing in dataset"

        obs_a = identical_header + tail_error_a
        obs_b = identical_header + tail_error_b

        hash_a = get_step_hash("sandbox_run", {"step": 2}, obs_a)
        hash_b = get_step_hash("sandbox_run", {"step": 2}, obs_b)

        assert hash_a != hash_b, "Identical headers masked distinct error tails!"

    def test_identical_tail_over_500_chars_produces_identical_hash(self):
        """Only the last 500 characters of the observation are hashed."""
        common_tail = "ERROR_TAIL_" + ("X" * 500)
        obs_1 = ("PrefixOne_" * 50) + common_tail
        obs_2 = ("PrefixTwo_" * 50) + common_tail

        assert obs_1[-500:] == obs_2[-500:]
        h1 = get_step_hash("tool", {"k": "v"}, obs_1)
        h2 = get_step_hash("tool", {"k": "v"}, obs_2)
        assert h1 == h2


# ===========================================================================
# 4. Tests for src.exporter
# ===========================================================================

class TestExporter:
    """Tests for run-flattening document templating and citation gates."""

    def test_run_flattening_empty_paragraphs_and_substitution(self, tmp_path):
        """Test placeholder tags [[KEY]] in normal, fragmented, and empty paragraphs."""
        template_file = tmp_path / "mock_template.docx"
        out_file = tmp_path / "output_memo.docx"

        # Create a mock .docx with varied paragraph structures
        doc = Document()
        # 1. Normal paragraph
        doc.add_paragraph("Header: [[ORGANIZATION]] - Ref: [[DOC_ID]]")
        # 2. Empty paragraph with no text or runs
        p_empty = doc.add_paragraph()
        assert len(p_empty.runs) == 0
        # 3. Fragmented paragraph where tag is split across multiple runs
        p_frag = doc.add_paragraph()
        p_frag.add_run("Component: [[")
        p_frag.add_run("EQUIPMENT")
        p_frag.add_run("]] Status")
        # 4. Plain paragraph without placeholders
        doc.add_paragraph("Standard narrative without tags.")

        doc.save(str(template_file))

        slots = {
            "ORGANIZATION": "MRPL Mangalore",
            "DOC_ID": "MEMO-2026-Q3",
            "EQUIPMENT": "Pump P-101A",
        }
        citations = ["[SOP-REF §3.2 p.14]"]

        result_path = render_deliverable(
            template=str(template_file),
            out=str(out_file),
            slots=slots,
            citations=citations,
        )

        assert result_path == str(out_file)
        assert out_file.is_file()

        # Verify output document contents
        res_doc = Document(str(out_file))
        assert len(res_doc.paragraphs) >= 4
        assert res_doc.paragraphs[0].text == "Header: MRPL Mangalore - Ref: MEMO-2026-Q3"
        assert res_doc.paragraphs[1].text == ""  # Empty paragraph was preserved without error
        assert res_doc.paragraphs[2].text == "Component: Pump P-101A Status"
        assert res_doc.paragraphs[3].text == "Standard narrative without tags."
        # Verify deliverable fallback content was appended because [[CONTENT]] was absent
        paragraphs_text = [p.text for p in res_doc.paragraphs]
        assert "Executive Summary & Technical Memo" in paragraphs_text
        assert "• [SOP-REF §3.2 p.14]" in paragraphs_text

    def test_citation_contract_validation(self, tmp_path):
        """Verify unresolvable or malformed citations are blocked by the exporter gate."""
        template_file = tmp_path / "test_tpl.docx"
        doc = Document()
        doc.add_paragraph("Test memo")
        doc.save(str(template_file))

        # Malformed citations must trigger assertion failure
        invalid_citations = [
            "SOP-REF §3.2 p.14",             # Missing brackets
            "[SOP-REF §3 p.14]",             # Single digit section (§3 vs §3.2)
            "[SOP-REF §3.2]",                # Missing page
            "[REF §3.2 p.14]",               # Missing SOP-
            "[SOP-REF §3.2 p.XIV]",          # Roman numeral page
        ]

        for bad_cite in invalid_citations:
            with pytest.raises(AssertionError, match="unresolvable citation blocked"):
                render_deliverable(
                    template=str(template_file),
                    out=str(tmp_path / "fail.docx"),
                    citations=[bad_cite],
                )

        # Valid citation must pass
        render_deliverable(
            template=str(template_file),
            out=str(tmp_path / "pass.docx"),
            citations=["[SOP-REF §3.2 p.14]", "[SOP-REF §10.1 p.5]"],
        )
        assert (tmp_path / "pass.docx").is_file()

    def test_render_deliverable_fallback_missing_template(self, tmp_path):
        """Verify fallback document initialization when template file is missing."""
        out_file = tmp_path / "fallback_output.docx"
        slots = {
            "TITLE": "Safety Review Memo",
            "DATE": "2026-09-02",
            "CONTENT": "Inspection completed successfully without deviations.",
            "CITATIONS": "[SOP-REF §3.2 p.14]",
        }

        result_path = render_deliverable(
            template=str(tmp_path / "non_existent_template.dotx"),
            out=str(out_file),
            slots=slots,
            citations=["[SOP-REF §3.2 p.14]"],
        )

        assert os.path.isabs(result_path)
        assert os.path.exists(result_path)

        res_doc = Document(result_path)
        # Verify heading and placeholders
        paragraphs_text = [p.text for p in res_doc.paragraphs]
        assert "SOVEREIGN WORKBENCH - ENGINEERING MEMORANDUM" in paragraphs_text[0]
        assert "Safety Review Memo" in paragraphs_text
        assert "2026-09-02" in paragraphs_text
        assert "Inspection completed successfully without deviations." in paragraphs_text
        assert "[SOP-REF §3.2 p.14]" in paragraphs_text

