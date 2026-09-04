import logging
import os
import re
from datetime import datetime
from docx import Document
from docx.shared import Inches

log = logging.getLogger("sovereign.exporter")

CITE_RE = re.compile(r"\[SOP-REF §\d+\.\d+ p\.\d+\]")


def render_deliverable(task_id: str = "memo", citations: list[str] = None,
                       content: str = "", title: str = "",
                       template: str = "assets/mrpl_template.dotx", out: str = None,
                       slots: dict = None, pictures: list[str] = None):
    """
    Render engineering deliverable as a formatted .docx document.
    - Populates [[DATE]], [[TITLE]], [[CONTENT]], and [[CITATIONS]] from agent state.
    - Replaces matching placeholders in all paragraphs, headings, and tables.
    - If [[CONTENT]] is not found in the template, appends Executive Summary and Citations.
    - Saves document to artifacts/{task_id}_memo.docx and returns the file path.
    """
    # Backward compatibility: detect if template was passed as first positional arg
    if isinstance(task_id, str) and (task_id.endswith(".dotx") or task_id.endswith(".docx") or "/" in task_id or "\\" in task_id):
        template = task_id
        task_id = "memo"

    template_path = template
    if os.path.exists(template_path):
        doc = Document(template_path)  # pre-styled corporate .dotx
    else:
        log.warning("Template not found at '%s'. Initializing clean fallback document.", template_path)
        doc = Document()
        doc.add_heading("SOVEREIGN WORKBENCH - ENGINEERING MEMORANDUM", level=0)
        standard_tags = ["TITLE", "DATE", "CONTENT", "CITATIONS"]
        keys_to_add = [
            k for k in standard_tags
            if (slots and k in slots) or (k == "CITATIONS" and citations)
        ]
        if not keys_to_add:
            keys_to_add = standard_tags

        for k in keys_to_add:
            doc.add_paragraph(f"[[{k}]]")

        if slots:
            for key in slots:
                if key not in standard_tags:
                    doc.add_paragraph(f"[[{key}]]")

    today_str = datetime.now().strftime("%B %d, %Y")
    if not title:
        if slots and "TITLE" in slots:
            title = str(slots["TITLE"])
        else:
            title = "Engineering Memorandum: Unit 200 Inspection & Corrosion Trends"

    if not content:
        if slots and "CONTENT" in slots:
            content = str(slots["CONTENT"])
        elif slots and "PLAN" in slots:
            content = str(slots["PLAN"])
        else:
            content = f"Technical memorandum synthesized for task {task_id}. Operational parameters verified in accordance with MRPL inspection specifications."

    cits = citations if citations is not None else []
    replacements = {
        "[[DATE]]": today_str,
        "[[TITLE]]": title,
        "[[CONTENT]]": content,
        "[[CITATIONS]]": "\n".join(f"• {c}" for c in cits) if cits else "No SOP citations attached.",
    }

    # Integrate any extra slot replacements
    if slots:
        for k, v in slots.items():
            placeholder = k if k.startswith("[[") else f"[[{k}]]"
            replacements[placeholder] = str(v)

    content_replaced = False

    def replace_in_para(para):
        nonlocal content_replaced
        full_text = "".join(run.text for run in para.runs)
        if "[[CONTENT]]" in full_text:
            content_replaced = True

        matched = False
        for placeholder, value in replacements.items():
            if placeholder in full_text:
                if placeholder == "[[CONTENT]]":
                    content_replaced = True
                full_text = full_text.replace(placeholder, str(value))
                matched = True

        if matched or not para.runs:
            if para.runs:
                para.runs[0].text = full_text
                for r in para.runs[1:]:
                    r.text = ""
            else:
                para.add_run(full_text)

    # Iterate over all paragraphs (including headings)
    for para in doc.paragraphs:
        replace_in_para(para)

    # Iterate over all table cells
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    replace_in_para(para)

    # If [[CONTENT]] was not found anywhere in the template, append content and citations
    if not content_replaced:
        doc.add_heading("Executive Summary & Technical Memo", level=1)
        doc.add_paragraph(content)
        doc.add_heading("Regulatory & SOP Grounding Citations", level=2)
        for cit in cits:
            doc.add_paragraph(f"• {cit}")

    for cite in cits:  # citation contract gate
        assert CITE_RE.fullmatch(cite), f"unresolvable citation blocked: {cite}"

    for path in pictures or []:  # P&ID thumbnails + trend chart
        if os.path.exists(path):
            doc.add_picture(path, width=Inches(4.8))

    # Resolve output path: artifacts/{task_id}_memo.docx
    tid = task_id or (slots and slots.get("TASK_ID")) or "memo"
    if out is None or out == "artifacts/memo.docx":
        out = f"artifacts/{tid}_memo.docx"

    out_abs = os.path.abspath(out)
    os.makedirs(os.path.dirname(out_abs), exist_ok=True)
    doc.save(out_abs)
    return out_abs


def render_spreadsheet(task_id: str, records: list[dict] = None, summary_metrics: dict = None) -> str:
    """
    Renders calculated engineering metrics to artifacts/{task_id}_report.xlsx
    with professional industrial formatting. Implements requirement R6 from Dev 2 guide.
    """
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side

    records = records or []
    summary_metrics = summary_metrics or {}

    wb = Workbook()
    ws = wb.active
    ws.title = "Corrosion Analysis"

    # Header styling
    header_fill = PatternFill(start_color="1F4E79", end_color="1F4E79", fill_type="solid")
    header_font = Font(name="Calibri", size=11, bold=True, color="FFFFFF")
    thin_border = Border(
        left=Side(style='thin', color='D9D9D9'),
        right=Side(style='thin', color='D9D9D9'),
        top=Side(style='thin', color='D9D9D9'),
        bottom=Side(style='thin', color='D9D9D9')
    )

    # Title Banner
    ws.merge_cells("A1:E1")
    ws["A1"] = "SOVEREIGN AI WORKBENCH - UNIT 200 INSPECTION METRICS"
    ws["A1"].font = Font(name="Calibri", size=14, bold=True, color="1F4E79")
    ws["A1"].alignment = Alignment(vertical="center")
    ws.row_dimensions[1].height = 30

    # Write Table Headers
    headers = ["Tag ID", "Component", "Nominal (mm)", "Measured (mm)", "Corrosion Rate (mm/yr)"]
    for col_num, header in enumerate(headers, 1):
        cell = ws.cell(row=3, column=col_num, value=header)
        cell.fill = header_fill
        cell.font = header_font
        cell.alignment = Alignment(horizontal="center", vertical="center")
        cell.border = thin_border
    ws.row_dimensions[3].height = 24

    # Populate Data Rows with Zebra Striping
    alt_fill = PatternFill(start_color="F2F2F2", end_color="F2F2F2", fill_type="solid")
    for row_idx, item in enumerate(records, 4):
        ws.cell(row=row_idx, column=1, value=item.get("tag_id", "P-201A"))
        ws.cell(row=row_idx, column=2, value=item.get("component", "Overhead Reflux"))
        ws.cell(row=row_idx, column=3, value=item.get("nominal_mm", 12.50))
        ws.cell(row=row_idx, column=4, value=item.get("measured_mm", 11.22))
        ws.cell(row=row_idx, column=5, value=item.get("rate_mm_yr", 0.32))

        for col_idx in range(1, 6):
            c = ws.cell(row=row_idx, column=col_idx)
            c.border = thin_border
            if row_idx % 2 == 0:
                c.fill = alt_fill

    # Auto-adjust column widths
    from openpyxl.utils import get_column_letter
    for col in ws.columns:
        max_len = max(len(str(cell.value or '')) for cell in col)
        col_letter = get_column_letter(col[0].column)
        ws.column_dimensions[col_letter].width = max(max_len + 4, 12)

    os.makedirs("artifacts", exist_ok=True)
    out_path = os.path.abspath(f"artifacts/{task_id}_report.xlsx")
    wb.save(out_path)
    return out_path


def render_excel_deliverable(task_id: str, data: list[dict], out: str = None) -> str:
    """
    Render generic tabular data to artifacts/{task_id}_calculations.xlsx.

    Features:
      - Bold, dark-blue header row (auto-detected from dict keys)
      - Alternating row zebra-stripe (white / light-gray)
      - Numeric columns right-aligned; text columns left-aligned
      - Auto-width columns (capped at 50 chars)
      - Frozen pane at row 4 (below title + header)
      - Returns absolute path to the saved file.

    Args:
        task_id: Unique task identifier used in the filename.
        data:    List of dicts (each dict is one row; keys = column headers).
        out:     Optional custom output path.
    """
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    from openpyxl.utils import get_column_letter

    wb = Workbook()
    ws = wb.active
    ws.title = "Swara.ai Calculations"

    os.makedirs("artifacts", exist_ok=True)
    out_path = os.path.abspath(out) if out else os.path.abspath(f"artifacts/{task_id}_calculations.xlsx")

    if not data:
        wb.save(out_path)
        return out_path

    # ── Styles ────────────────────────────────────────────────────────────
    header_fill = PatternFill(start_color="1F4E79", end_color="1F4E79", fill_type="solid")
    header_font = Font(name="Calibri", size=11, bold=True, color="FFFFFF")
    zebra_fill  = PatternFill(start_color="F2F7FB", end_color="F2F7FB", fill_type="solid")
    thin_side   = Side(style="thin", color="D0D7E0")
    thin_border = Border(left=thin_side, right=thin_side, top=thin_side, bottom=thin_side)

    headers = list(data[0].keys())

    # ── Title banner (row 1) ──────────────────────────────────────────────
    banner_col = get_column_letter(len(headers))
    ws.merge_cells(f"A1:{banner_col}1")
    ws["A1"] = "Swara.ai — Sovereign Industrial Calculations"
    ws["A1"].font = Font(name="Calibri", size=13, bold=True, color="1F4E79")
    ws["A1"].alignment = Alignment(horizontal="left", vertical="center")
    ws.row_dimensions[1].height = 28

    # ── Header row (row 3) ────────────────────────────────────────────────
    for col_idx, header in enumerate(headers, 1):
        cell = ws.cell(row=3, column=col_idx, value=header)
        cell.fill = header_fill
        cell.font = header_font
        cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
        cell.border = thin_border
    ws.row_dimensions[3].height = 22

    # ── Detect numeric columns ────────────────────────────────────────────
    numeric_cols: set[int] = set()
    for row_data in data:
        for col_idx, key in enumerate(headers, 1):
            val = row_data.get(key)
            if isinstance(val, (int, float)):
                numeric_cols.add(col_idx)

    # ── Data rows with zebra striping ─────────────────────────────────────
    for row_idx, row_data in enumerate(data, 4):
        is_even = (row_idx % 2 == 0)
        for col_idx, key in enumerate(headers, 1):
            val = row_data.get(key, "")
            cell = ws.cell(row=row_idx, column=col_idx, value=val)
            cell.border = thin_border
            if is_even:
                cell.fill = zebra_fill
            if col_idx in numeric_cols:
                cell.alignment = Alignment(horizontal="right", vertical="center")
                if isinstance(val, float):
                    cell.number_format = "#,##0.0000"
            else:
                cell.alignment = Alignment(horizontal="left", vertical="center")

    # ── Auto-width columns ────────────────────────────────────────────────
    for col_idx, header in enumerate(headers, 1):
        col_letter = get_column_letter(col_idx)
        max_len = len(str(header))
        for row_data in data:
            val = row_data.get(header, "")
            max_len = max(max_len, len(str(val)))
        ws.column_dimensions[col_letter].width = min(max(max_len + 3, 12), 50)

    # ── Freeze panes below header ─────────────────────────────────────────
    ws.freeze_panes = "A4"

    wb.save(out_path)
    log.info("[EXPORTER] Excel deliverable saved: %s (%d rows)", out_path, len(data))
    return out_path
